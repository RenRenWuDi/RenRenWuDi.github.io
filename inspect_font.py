from docx import Document

doc = Document(r'E:\区块链编程文章\12-1-DeFi风险管理.docx')
out = []
# Inspect fonts on various paragraphs
for idx in [6, 11, 12, 13, 14, 15, 27, 50]:
    if idx >= len(doc.paragraphs):
        continue
    p = doc.paragraphs[idx]
    fonts = set()
    szs = set()
    for r in p.runs:
        f = r.font.name
        if f:
            fonts.add(f)
        if r.font.size:
            szs.add(str(r.font.size))
    out.append(f'[{idx}] fonts={fonts} sizes={szs} | text={p.text[:35]!r}')

# Count paragraphs by dominant font
from collections import Counter
font_counter = Counter()
for p in doc.paragraphs:
    fonts = set(r.font.name for r in p.runs if r.font.name)
    if fonts:
        font_counter[str(sorted(fonts))] += 1
out.append('')
out.append('=== Font distribution across paragraphs ===')
for k, v in font_counter.most_common(10):
    out.append(f'  {k}: {v}')

with open(r'F:\tech-blog\inspect_font.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(out))
print('written', len(out), 'lines')
