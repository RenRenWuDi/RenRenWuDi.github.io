from docx import Document
from collections import Counter

doc = Document(r'E:\区块链编程文章\已发\2-1-Solidity基础语法.docx')
fc = Counter()
for p in doc.paragraphs:
    fonts = set(r.font.name for r in p.runs if r.font.name)
    if fonts:
        fc[str(sorted(fonts))] += 1
print('2-1 font dist:', dict(fc))
for i, p in enumerate(doc.paragraphs[:8]):
    s = p.style.name if p.style else ''
    fns = set(r.font.name for r in p.runs if r.font.name)
    print(f'[{i}] style={s} font={fns} | {p.text[:40]!r}')
