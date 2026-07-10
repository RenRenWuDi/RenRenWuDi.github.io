from docx import Document

doc = Document(r'E:\区块链编程文章\12-1-DeFi风险管理.docx')
out = []
out.append('=== PARAGRAPHS (style | text preview) ===')
for i, p in enumerate(doc.paragraphs[:45]):
    s = p.style.name if p.style else ''
    t = p.text.replace('\n', ' ').strip()
    out.append(f'[{i:02d}] {s:22s} | {t[:75]}')
out.append('')
out.append(f'=== total paragraphs: {len(doc.paragraphs)}')

# Look at code-like paragraphs
out.append('')
out.append('=== CODE DETECTION (paragraphs with monospace/code style) ===')
code_styles = set()
for p in doc.paragraphs:
    s = p.style.name if p.style else ''
    if s and ('Code' in s or 'code' in s or 'Source' in s):
        code_styles.add(s)
out.append(f'Code-related styles found: {code_styles}')

# Sample 3 code paragraphs
cnt = 0
for i, p in enumerate(doc.paragraphs):
    s = p.style.name if p.style else ''
    if s and ('Code' in s or 'code' in s or 'Source' in s):
        out.append(f'  CODE[{i}] style={s} | {p.text[:90]}')
        cnt += 1
        if cnt >= 5:
            break

with open(r'F:\tech-blog\inspect_12_1.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(out))
print('written')
